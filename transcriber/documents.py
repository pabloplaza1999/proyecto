"""Generación de actas organizadas en .docx y .txt."""

from __future__ import annotations

import logging
import os
import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


def _format_time(seconds: float) -> str:
    """Convierte segundos a formato [HH:MM:SS]."""
    td = datetime.timedelta(seconds=int(seconds))
    return f"[{str(td)}]"


def _group_segments_into_paragraphs(segments: list[dict], pause_threshold: float = 1.5) -> list[dict]:
    """
    Agrupa frases cortas en párrafos más grandes si no hubo mucho silencio entre ellas.
    Retorna una lista de párrafos con su timestamp inicial.
    """
    if not segments:
        return []

    paragraphs = []
    current_paragraph = {
        "start": segments[0]["start"],
        "text": segments[0]["text"]
    }
    
    for i in range(1, len(segments)):
        prev_segment = segments[i - 1]
        curr_segment = segments[i]
        
        # Calcula cuánto silencio hubo entre la frase anterior y la actual
        silence_gap = curr_segment["start"] - prev_segment["end"]
        
        # Si el silencio es mayor al límite, asumimos que es otro orador u otro tema
        if silence_gap >= pause_threshold:
            paragraphs.append(current_paragraph)
            current_paragraph = {
                "start": curr_segment["start"],
                "text": curr_segment["text"]
            }
        else:
            # Si hablaron de corrido, unimos la frase al párrafo actual
            current_paragraph["text"] += " " + curr_segment["text"]
            
    # Agrega el último párrafo
    paragraphs.append(current_paragraph)
    return paragraphs


def create_docx(segments: list[dict], output_path: str, video_name: str) -> bool:
    """Crea un documento Word organizado, con título, metadata y timestamps."""
    try:
        doc = Document()
        
        # --- ENCABEZADO Y TÍTULO ---
        title = doc.add_heading("ACTA DE TRANSCRIPCIÓN", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        meta = doc.add_paragraph()
        meta.add_run("Archivo original: ").bold = True
        meta.add_run(video_name + "\n")
        meta.add_run("Fecha de generación: ").bold = True
        meta.add_run(datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S") + "\n")
        doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # --- CUERPO DE LA TRANSCRIPCIÓN ---
        if not segments:
            doc.add_paragraph("No se detectó audio en este video.")
        else:
            paragraphs = _group_segments_into_paragraphs(segments)
            
            for p in paragraphs:
                timestamp = _format_time(p["start"])
                
                # Crear párrafo
                para = doc.add_paragraph()
                
                # Agregar el tiempo en color gris y negrita
                time_run = para.add_run(f"{timestamp} ")
                time_run.bold = True
                time_run.font.color.rgb = RGBColor(100, 100, 100) # Gris oscuro
                
                # Agregar el texto
                para.add_run(p["text"])
                
                # Espacio extra entre párrafos
                para.paragraph_format.space_after = Pt(12)

        doc.save(output_path)
        file_size = os.path.getsize(output_path) / 1024
        logger.info("Word creado: %s (%.2f KB)", output_path, file_size)
        return True
    except Exception as error:
        logger.error("Error creando Word: %s", error)
        return False


def create_txt(segments: list[dict], output_path: str, video_name: str) -> bool:
    """Crea un archivo de texto plano limpio y ordenado."""
    try:
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("=" * 60 + "\n")
            f.write(f" ACTA DE TRANSCRIPCIÓN - {video_name}\n")
            f.write(f" Fecha: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("=" * 60 + "\n\n")
            
            if not segments:
                f.write("No se detectó audio en este video.\n")
            else:
                paragraphs = _group_segments_into_paragraphs(segments)
                for p in paragraphs:
                    timestamp = _format_time(p["start"])
                    f.write(f"{timestamp}\n{p['text']}\n\n")
                    
        file_size = os.path.getsize(output_path) / 1024
        logger.info("TXT creado: %s (%.2f KB)", output_path, file_size)
        return True
    except Exception as error:
        logger.error("Error creando TXT: %s", error)
        return False
